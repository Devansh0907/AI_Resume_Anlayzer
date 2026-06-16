import os
import re
import json
import io
import uvicorn
from fastapi import FastAPI, File, UploadFile, Form, Header, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pypdf import PdfReader
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import google.generativeai as genai

app = FastAPI(title="Resume Analyzer AI & NLP Service")

# Allow CORS for integration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load spaCy model with safety fallback
nlp = None
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
except Exception as e:
    print(f"Warning: Failed to load spaCy model en_core_web_sm: {e}")
    print("Using regex-based tokenization fallback.")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse PDF file: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse DOCX file: {str(e)}")

def clean_text(text: str) -> str:
    # Basic cleaning
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def tokenize_nlp(text: str):
    """Tokenize and clean text using spaCy if available, else simple regex fallback."""
    cleaned = clean_text(text)
    if nlp is not None:
        try:
            doc = nlp(cleaned)
            # Filter stop words, punctuation, and return lemmatized forms
            tokens = [token.lemma_ for token in doc if not token.is_stop and not token.is_punct and len(token.text.strip()) > 1]
            return tokens
        except Exception:
            pass
    # Regex fallback
    tokens = re.findall(r'\b[a-zA-Z0-9]{2,}\b', cleaned)
    # Basic stop words list to filter
    stop_words = {"the", "and", "a", "of", "to", "in", "is", "for", "with", "on", "as", "by", "at", "an", "this", "that", "from", "it", "are", "be", "or"}
    return [t for t in tokens if t not in stop_words]

def compute_cosine_similarity(text1: str, text2: str) -> float:
    try:
        # Preprocess text by tokenizing and joining
        tokens1 = " ".join(tokenize_nlp(text1))
        tokens2 = " ".join(tokenize_nlp(text2))
        
        if not tokens1 or not tokens2:
            return 0.0
            
        vectorizer = TfidfVectorizer()
        tfidf = vectorizer.fit_transform([tokens1, tokens2])
        sim = cosine_similarity(tfidf[0:1], tfidf[1:2])
        return float(sim[0][0]) * 100.0
    except Exception as e:
        print(f"Error computing cosine similarity: {e}")
        return 0.0

def extract_nlp_keywords(text: str, top_n=20):
    """Extract key technical terms or nouns from text."""
    cleaned = clean_text(text)
    if nlp is not None:
        try:
            doc = nlp(cleaned)
            # Focus on Nouns, Proper Nouns, and Adjectives as potential keywords
            keywords = [token.text for token in doc if token.pos_ in {"NOUN", "PROPN", "ADJ"} and not token.is_stop and len(token.text) > 2]
            # Get frequency map
            freq = {}
            for kw in keywords:
                freq[kw] = freq.get(kw, 0) + 1
            sorted_kws = sorted(freq.items(), key=lambda x: x[1], reverse=True)
            return [kw[0] for kw in sorted_kws[:top_n]]
        except Exception:
            pass
    
    # Regex fallback
    tokens = re.findall(r'\b[a-zA-Z0-9]{3,}\b', cleaned)
    stop_words = {"the", "and", "that", "with", "this", "from", "for", "your", "will", "have", "are", "our", "you"}
    keywords = [t for t in tokens if t not in stop_words]
    freq = {}
    for kw in keywords:
        freq[kw] = freq.get(kw, 0) + 1
    sorted_kws = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return [kw[0] for kw in sorted_kws[:top_n]]

@app.post("/analyze")
async def analyze(
    file: UploadFile = File(...),
    jobDescription: str = Form(...),
    x_gemini_key: str = Header(None)
):
    # Determine API key priority: Header -> Env Var
    api_key = x_gemini_key or os.environ.get("GEMINI_API_KEY")
    
    if not api_key:
        raise HTTPException(
            status_code=400,
            detail="Gemini API Key is missing. Please configure GEMINI_API_KEY environment variable or provide it in the X-Gemini-Key header."
        )

    # 1. Read and parse file
    file_bytes = await file.read()
    filename = file.filename.lower()
    
    if filename.endswith(".pdf"):
        resume_text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        resume_text = extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF or DOCX.")

    if not resume_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract any text from the uploaded resume file.")

    # 2. NLP Cosine Similarity Analysis
    nlp_similarity = compute_cosine_similarity(resume_text, jobDescription)
    
    # Extract keywords
    jd_keywords = extract_nlp_keywords(jobDescription, top_n=25)
    resume_keywords = set(extract_nlp_keywords(resume_text, top_n=100))
    
    matched_nlp_skills = [kw for kw in jd_keywords if kw in resume_keywords]
    missing_nlp_skills = [kw for kw in jd_keywords if kw not in resume_keywords]

    # 3. LLM Deep AI Analysis
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.5-flash")
        
        prompt = f"""You are an advanced AI Applicant Tracking System (ATS) auditor.
Analyze the resume content against the job description provided.
Your assessment must be returned in strict JSON format matching the schema below.
Ensure the response is raw JSON without markdown formatting (do not wrap in ```json or ```).

Required JSON format:
{{
  "atsScore": <integer, overall match score from 0 to 100 based on core requirements alignment>,
  "matchPercentage": <integer, percentage score representing resume suitability for the role (0 to 100)>,
  "matchedSkills": [<string>, list of specific matching technologies, skills, or methodologies found in the resume],
  "missingSkills": [<string>, list of critical skills or technologies requested in the JD but missing or weak in the resume],
  "strengths": [<string>, list of 3-4 key advantages / highlights of the resume matching the JD],
  "weaknesses": [<string>, list of 3-4 areas of misalignment, gaps, or structural flaws],
  "suggestions": [<string>, list of 4-5 concrete, actionable steps to improve the resume for this specific job]
}}

Resume:
{resume_text}

Job Description:
{jobDescription}
"""

        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Clean potential markdown wrappers if the LLM output includes them
        if response_text.startswith("```"):
            # Strip first line starting with ```
            response_text = re.sub(r'^```[a-zA-Z]*\n', '', response_text)
            # Strip trailing ```
            response_text = re.sub(r'\n```$', '', response_text)
            response_text = response_text.strip()
            
        ai_data = json.loads(response_text)
        
    except json.JSONDecodeError as e:
        print(f"Error parsing Gemini JSON: {e}. Raw text: {response_text}")
        # Build a structured fallback parsing from raw text or mock values if JSON fails
        ai_data = {
            "atsScore": int((nlp_similarity + 50) / 2),
            "matchPercentage": int(nlp_similarity),
            "matchedSkills": matched_nlp_skills[:10],
            "missingSkills": missing_nlp_skills[:10],
            "strengths": ["Parsed resume text matches the structure", "NLP parsing detected key term alignment"],
            "weaknesses": ["Gemini JSON parsing failure, formatting fallback active"],
            "suggestions": ["Ensure resume formatting uses standard PDF/Word headers"]
        }
    except Exception as e:
        print(f"Error calling Gemini API: {e}")
        # General fallback if API fails
        ai_data = {
            "atsScore": int(nlp_similarity),
            "matchPercentage": int(nlp_similarity),
            "matchedSkills": matched_nlp_skills[:10],
            "missingSkills": missing_nlp_skills[:10],
            "strengths": ["Successfully extracted text content", "NLP text matching active"],
            "weaknesses": ["AI Deep Matching is currently unavailable (API call failed)"],
            "suggestions": ["Please verify your Gemini API key is valid", "Try checking service logs for Gemini API connectivity"]
        }

    # 4. Integrate NLP Cosine score and keywords with AI data
    # We want to provide a highly comprehensive combined analysis response.
    # We will append the NLP similarity score and NLP-based keywords for display.
    combined_response = {
        "atsScore": ai_data.get("atsScore", 70),
        "matchPercentage": ai_data.get("matchPercentage", 70),
        # Merge lists to prevent empty results, keeping unique values
        "matchedSkills": list(set(ai_data.get("matchedSkills", []) + matched_nlp_skills)),
        "missingSkills": list(set(ai_data.get("missingSkills", []) + missing_nlp_skills)),
        "strengths": ai_data.get("strengths", []),
        "weaknesses": ai_data.get("weaknesses", []),
        "suggestions": ai_data.get("suggestions", []),
        "nlpCosineSimilarity": round(nlp_similarity, 1)
    }

    return combined_response

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
